import os
# Suppress macOS deprecation warnings from NSOpenPanel in Tkinter
os.environ['TK_SILENCE_DEPRECATION'] = '1'

import csv
import time
import json
import subprocess
import requests
import numpy as np
import pandas as pd
import tkinter as tk
from tkinter import ttk, messagebox, filedialog, scrolledtext
from docx import Document
from typing import List, Dict
import re
import zipfile
import webbrowser

# Qdrant and LangChain imports
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings.base import Embeddings
from langchain_qdrant import Qdrant  # Make sure you have updated this package
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, VectorParams, Filter, FieldCondition, MatchValue
from langchain_community.vectorstores import Qdrant as QdrantVectorStore

######################################################
# GLOBAL CONSTANTS / DEFAULTS
######################################################
EMBEDDING_API_ENDPOINT = "https://apim.stanfordhealthcare.org/openai3/deployments/text-embedding-3-large/embeddings?api-version=2023-05-15"
RATE_LIMIT_TOKENS_PER_MINUTE = 15000
tokens_used = 0
start_time = time.time()

DEFAULT_DOWNLOADS = os.path.join(os.path.expanduser("~"), "Downloads")

######################################################
# SCROLLABLE FRAME HELPER
######################################################
class ScrollableFrame(ttk.Frame):
    def __init__(self, parent, *args, **kwargs):
        super().__init__(parent, *args, **kwargs)

        # The canvas widget.
        self.canvas = tk.Canvas(self)
        
        # Vertical Scrollbar
        self.scrollbar_y = ttk.Scrollbar(self, orient="vertical", command=self.canvas.yview)
        self.scrollbar_y.pack(side="right", fill="y")
        
        # Horizontal Scrollbar
        self.scrollbar_x = ttk.Scrollbar(self, orient="horizontal", command=self.canvas.xview)
        self.scrollbar_x.pack(side="bottom", fill="x")
        
        # Frame inside the canvas that will actually hold your widgets
        self.scrollable_frame = ttk.Frame(self.canvas)
        
        # Create a window within the canvas to place the scrollable frame
        self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        self.canvas.pack(side="left", fill="both", expand=True)

        # Connect the scrollbars
        self.canvas.configure(
            yscrollcommand=self.scrollbar_y.set,
            xscrollcommand=self.scrollbar_x.set
        )

        # Update the scroll region whenever the scrollable_frame changes size
        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all"))
        )

######################################################
# DOCX READER
######################################################
# def read_docx_file(docx_path: str) -> str:
#     doc = Document(docx_path)
#     return "\n".join(p.text for p in doc.paragraphs)
def read_docx_file(docx_path: str) -> str:
    doc = Document(docx_path)
    texts = []
    # Extract text from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text.strip())
    # Also extract text from tables (if any)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    texts.append(cell.text.strip())
    return "\n".join(texts)


######################################################
# EMBEDDINGS CLASS
######################################################
class StanfordSecureEmbeddings(Embeddings):
    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        global tokens_used, start_time
        embeddings = []
        max_retries = 5
        backoff_factor = 2
        initial_wait = 1
        batch_size = 16
        for i in range(0, len(texts), batch_size):
            batch_texts = texts[i: i + batch_size]
            tokens_needed = sum(len(txt.split()) for txt in batch_texts)
            elapsed_time = time.time() - start_time
            if tokens_used + tokens_needed > RATE_LIMIT_TOKENS_PER_MINUTE:
                wait_time = 60 - elapsed_time
                if wait_time > 0:
                    print(f"[Embeddings] Rate limit reached. Waiting {wait_time:.2f} seconds...")
                    time.sleep(wait_time)
                tokens_used = 0
                start_time = time.time()
            headers = {"Ocp-Apim-Subscription-Key": App.api_key, "Content-Type": "application/json"}
            payload = {"input": batch_texts}
            retries = 0
            while retries < max_retries:
                try:
                    response = requests.post(EMBEDDING_API_ENDPOINT, headers=headers, json=payload)
                    if response.status_code == 429:
                        retry_after = int(response.headers.get("Retry-After", 0))
                        wait_time = retry_after if retry_after > 0 else initial_wait * (backoff_factor ** retries)
                        print(f"[Embeddings] 429 received. Sleeping {wait_time}s before retrying...")
                        time.sleep(wait_time)
                        retries += 1
                        continue
                    response.raise_for_status()
                    result = response.json()
                    tokens_used += tokens_needed
                    batch_embeddings = [data["embedding"] for data in result["data"]]
                    embeddings.extend(batch_embeddings)
                    break
                except Exception as e:
                    print(f"[Embeddings] Error: {e}")
                    retries += 1
                    time.sleep(initial_wait * (backoff_factor ** retries))
            else:
                print("[Embeddings] Max retries exceeded. Using zero vectors.")
                embeddings.extend([np.zeros(3072).tolist()] * len(batch_texts))
        return embeddings

    def embed_query(self, text: str) -> List[float]:
        return self.embed_documents([text])[0]

######################################################
# LLM CLASSES
######################################################
class StanfordSecureGPT:
    API_ENDPOINT_GPT4O = "https://apim.stanfordhealthcare.org/openai3/deployments/gpt-4o/chat/completions?api-version=2023-05-15"
    API_ENDPOINT_O1MINI = "https://apim.stanfordhealthcare.org/openai3/deployments/o1-mini/chat/completions?api-version=2023-05-15"
    API_ENDPOINT_O1 = "https://apim.stanfordhealthcare.org/openai-eastus2/deployments/o1/chat/completions?api-version=2024-12-01-preview"
    def __init__(self, temperature=0.0, max_tokens=2000, model_selection="gpt4o"):
        self.temperature = temperature
        self.max_tokens = max_tokens
        self.model_selection = model_selection
    def invoke(self, prompt: str) -> str:
        headers = {"Ocp-Apim-Subscription-Key": App.api_key, "Content-Type": "application/json"}
        if self.model_selection == "gpt4o":
            payload = {"model": "gpt-4o", "messages": [{"role": "user", "content": prompt}],
                       "max_tokens": self.max_tokens, "temperature": self.temperature}
            endpoint = self.API_ENDPOINT_GPT4O
        elif self.model_selection == "o1-mini":
            payload = {"model": "o1-mini", "messages": [{"role": "user", "content": prompt}]}
            endpoint = self.API_ENDPOINT_O1MINI
        elif self.model_selection == "o1":
            payload = {"model": "o1", "messages": [{"role": "user", "content": prompt}]}
            endpoint = self.API_ENDPOINT_O1
        else:
            return "Unsupported model selection."
        max_retries = 5
        for attempt in range(max_retries):
            try:
                response = requests.post(endpoint, headers=headers, json=payload)
                if response.status_code == 429:
                    wait_time = int(response.headers.get("Retry-After", 1))
                    time.sleep(wait_time)
                    continue
                response.raise_for_status()
                result = response.json()
                return result["choices"][0]["message"]["content"].strip()
            except Exception as e:
                print(f"[LLM] Error (attempt {attempt+1}): {e}")
                time.sleep(2 ** attempt)
        return "Error: LLM API call failed."

class DeepSeekR1Client:
    API_ENDPOINT = "https://apim.stanfordhealthcare.org/deepseekr1/v1/chat/completions"
    def __init__(self, temperature=0.1, max_tokens=2000):
        self.temperature = temperature
        self.max_tokens = max_tokens
    def invoke(self, prompt: str) -> str:
        delay = 1
        max_retries = 5
        headers = {"Ocp-Apim-Subscription-Key": App.api_key, "Content-Type": "application/json"}
        payload = {"model": "deepseek-chat", "messages": [{"role": "user", "content": prompt}],
                   "temperature": self.temperature, "max_tokens": self.max_tokens, "top_p": 1, "stream": False}
        for attempt in range(max_retries):
            try:
                response = requests.post(self.API_ENDPOINT, headers=headers, json=payload)
                if response.status_code == 429:
                    time.sleep(delay)
                    delay *= 2
                    continue
                response.raise_for_status()
                result = response.json()
                return result["choices"][0]["message"]["content"].strip()
            except Exception as e:
                print(f"[DeepSeek] Error (attempt {attempt+1}): {e}")
                time.sleep(delay)
                delay *= 2
        return "Error: DeepSeek R1 API call failed."

######################################################
# MAIN APP WITH 4 TABS
######################################################
class App(tk.Tk):
    # Shared data
    api_key = ""
    qdrant_folder = ""
    vector_store = None

    # For "create new" steps
    docx_filenames = []  # list of DOCX file base names (without .docx)
    metadata_types = []  # list of metadata columns
    csv_data = None
    docx_folder = ""
    csv_template_path = ""

    # For storing document metadata (for Query and Analysis Grid filtering)
    all_doc_metadata = []

    def __init__(self):
        super().__init__()
        self.title("Stanford Secure GPT Qualitative Analysis Software")
        self.geometry("1400x900")
        self.resizable(True, True)
        self.notebook = ttk.Notebook(self)
        self.notebook.pack(fill="both", expand=True)
        self.welcome_page = WelcomePage(self.notebook, self)
        self.notebook.add(self.welcome_page, text="Welcome")
        self.api_page = APIKeyPage(self.notebook, self)
        self.notebook.add(self.api_page, text="Get Started")
        self.setup_page = SetupPage(self.notebook, self)
        self.notebook.add(self.setup_page, text="Setup")
        self.query_page = QueryTab(self.notebook, self)
        self.notebook.add(self.query_page, text="Query")
        self.analysis_grid_page = AnalysisGridTab(self.notebook, self)
        self.notebook.add(self.analysis_grid_page, text="Analysis Grid")
        # Initially disable Setup, Query, and Analysis Grid until API is verified
        self.notebook.tab(2, state="disabled")
        self.notebook.tab(3, state="disabled")
        self.notebook.tab(4, state="disabled")

######################################################
# PAGE 0: WELCOME
######################################################
class WelcomePage(ttk.Frame):
    def __init__(self, parent, controller: App):
        super().__init__(parent)
        self.controller = controller
        label = ttk.Label(self, text="Welcome to Stanford SecureGPT Qualitative Analysis Software", font=("Helvetica", 16))
        label.pack(pady=10)
        # Create a text widget to display the welcome message
        # dont want scroll text just want the text to be written directly on the screen
        welcome_text = tk.Text(self, height=100, width=160, wrap=tk.WORD)
        welcome_text.pack(pady=20, padx=20)
        
        # Insert the welcome message with formatting
        # make text bigger and easier to read
        welcome_text.config(font=("Helvetica", 18))
        welcome_message = """This tool is designed with two main capabilities:

1. Support for large, metadata-rich datasets
Most LLM interfaces limit how many documents you can upload at once. Our tool removes that barrier—you can upload entire folders of data. During upload, you can add metadata tags to each document (e.g., site name, interviewee ID, date), enabling you to filter and organize your dataset before querying. This gives you control over what subset of your data is used in any analysis.

2. Retrieval-Augmented Generation (RAG)
When working with large datasets, LLMs often struggle to consider all relevant information evenly—some sections get over-weighted, others ignored. 

To solve this, we use RAG. Here's how RAG works:

1. When you upload your documents, each paragraph or section is turned into a vector –a list of numbers that captures the meaning of the text.
2. When you ask a question, it's also turned into a vector that captures what you're asking about. 
3. The system compares your question vector to all the document vectors and finds the ones that are most similar in meaning.
4. Only those relevant excerpts are sent to the AI, so it can give a focused, accurate answer based on the most useful parts of your data.

Our goal is to facilitate easy access and use of the newest LLM models and most innovative computer science techniques for non-computer scientist researchers. Through using this application, you will familiarize yourself with Retrieval Augmented Generation and how to best use LLMs for your work. We hope this promotes collaboration between technical and non-technical researchers to improve AI tools and techniques for qualitative research applications."""

        welcome_text.insert(tk.END, welcome_message)
        welcome_text.configure(state='disabled')  # Make text read-only




######################################################
# PAGE 1: API KEY VERIFICATION
######################################################
class APIKeyPage(ttk.Frame):
    # i want to add this text to the page -- Before using this software, you'll need to enter your API key. This key is required to access the large language models through Stanford SecureGPT. Each research team has its own unique API key. Do not share your key with others.

# To request an API key, email Vicky Zhou <fzhou@stanfordhealthcare.org> with an explanation of what you are using the API key for. Note that SecureGPT is only available for research teams working with high-risk medical data. 

# An API (Application Programming Interface) lets one software tool communicate with another. Your API key is like a personalized password that identifies you when this tool talks to SecureGPT.

# While using the software, be sure you are on Stanford's campus and connected to Stanford's internet network or connected to Stanford's VPN ([instructions here](https://uit.stanford.edu/service/vpn)). 

    def __init__(self, parent, controller: App):
        super().__init__(parent)
        self.controller = controller
        api_text = tk.Text(self, height=12, width=160, wrap=tk.WORD)
        api_text.pack(pady=10, padx=10)
        api_text.config(font=("Helvetica", 18))
        api_message = """Before using this software, you'll need to enter your API key. This key is required to access the large language models through Stanford SecureGPT. Each research team has its own unique API key. Do not share your key with others.

To request an API key, email Vicky Zhou <fzhou@stanfordhealthcare.org> with an explanation of what you are using the API key for. Note that SecureGPT is only available for research teams working with high-risk medical data. 

An API (Application Programming Interface) lets one software tool communicate with another. Your API key is like a personalized password that identifies you when this tool talks to SecureGPT.

While using the software, be sure you are on Stanford's campus and connected to Stanford's internet network or connected to Stanford's VPN. VPN instructions can be found at https://uit.stanford.edu/service/vpn."""

        api_text.insert(tk.INSERT, api_message)
        api_text.configure(state='disabled') 
        label = ttk.Label(self, text="Enter and Verify Your SecureGPT API Key", font=("Helvetica", 16))
        label.pack(pady=10)
        ttk.Label(self, text="API Key:").pack(anchor="w", padx=5, pady=2)
        self.api_key_entry = ttk.Entry(self, width=60, show="*")
        self.api_key_entry.pack(anchor="w", padx=5, pady=2)
        self.verify_button = ttk.Button(self, text="Verify API Key", command=self.verify_api_key)
        self.verify_button.pack(pady=5)
        self.response_box = scrolledtext.ScrolledText(self, height=8)
        self.response_box.pack(fill="both", expand=True, padx=5, pady=5)
        # Make text read-only

    def verify_api_key(self):
        key = self.api_key_entry.get().strip()
        if not key:
            messagebox.showwarning("Input Error", "Please enter an API key.")
            return
        App.api_key = key
        self.response_box.delete("1.0", tk.END)
        self.response_box.insert(tk.END, "Verifying with GPT-4o...\n")
        test_llm = StanfordSecureGPT(temperature=0.0, max_tokens=50, model_selection="gpt4o")
        try:
            response = test_llm.invoke("Hello from the verification step!")
            self.response_box.insert(tk.END, f"Response:\n{response}\n")
            if "error" not in response.lower():
                messagebox.showinfo("Success", "API Key verified successfully.")
                self.controller.notebook.tab(2, state="normal")
                self.controller.notebook.tab(3, state="normal")
                self.controller.notebook.tab(4, state="normal")
                self.controller.notebook.select(1)
            else:
                self.response_box.insert(tk.END, "Verification might have failed. Check your key.\n")
        except Exception as e:
            self.response_box.insert(tk.END, f"Verification failed: {e}\n")

######################################################
# PAGE 2: SETUP (LOAD OR CREATE DB)
######################################################
class SetupPage(ttk.Frame):
    def __init__(self, parent, controller: App):
        super().__init__(parent)
        self.controller = controller

        # ---------------------------------------------------------------------
        # Outer container that splits the page into a left Frame (main UI)
        # and a right Frame (instructions sidebar).
        # ---------------------------------------------------------------------
        outer_frame = ttk.Frame(self)
        outer_frame.pack(fill="both", expand=True)

        # Left side for the existing Setup UI
        self.left_frame = ttk.Frame(outer_frame)
        # Expand to fill any leftover space; pinned to the left
        self.left_frame.pack(side="left", fill="both", expand=True, padx=5, pady=5)

        # Right side for the scrollable instructions
        right_frame = ttk.Frame(outer_frame, width=400)
        right_frame.pack(side="right", fill="y", padx=5, pady=5)

        # ---------------------------------------------------------------------
        # Add a label + scrolled text to show the instructions
        # ---------------------------------------------------------------------
        instructions_label = ttk.Label(right_frame, text="Setup Instructions", font=("Helvetica", 14, "bold"))
        instructions_label.pack(pady=5)

        instructions_text = scrolledtext.ScrolledText(right_frame, width=40, wrap="word")
        # Make the text a bit more readable: larger font
        instructions_text.config(font=("Helvetica", 15))
        instructions_text.pack(fill="both", expand=True, padx=5, pady=(0,5))

        # Insert your instructions text with line breaks for easier reading
        instructions_text.insert("1.0", """In the setup, we take your textual data and convert it into numerical vectors, which enables Retrieval-Augmented Generation (RAG). To generate these vectors, we use OpenAI's text-3-embedding-large embedding model. You can learn more about embeddings and this specific model here. Essentially, text-3-embedding-large transforms natural language into high-dimensional vector representations that capture the semantic meaning of your text. This allows similar pieces of text to be located close to each other in vector space, even if they don't share exact words, enabling more contextually relevant responses when used in combination with language models.

Once we generate these embeddings, we store them in a vector database called Qdrant DB. Qdrant is designed specifically for handling large collections of high-dimensional vectors and allows for efficient similarity search. While there are several vector databases available (such as Pinecone or FAISS), we use Qdrant because it supports pre-search metadata filtering. This means that when querying your dataset, we can first limit the scope of the search to only a specific subset of your data (e.g., by date, topic, or interview site), and then retrieve the most relevant items from within that subset. This is more efficient and precise than post-search filtering, which searches the entire dataset first and only filters the results afterward. You can learn more about Qdrant here.

Note that this process happens entirely on your machine. The database is created and stored on your computer and is not uploaded to the cloud or shared externally—unless you choose to do so.

You will see options to load or create Qdrant DB.

Create DB Option: If this is first time use with a new dataset, follow these steps:
Select Create DB option. 
Select Docx folder - Upload a folder with all of the data in .docx format. This folder can only have the .docx files, no subfolders or other file types are allowed.
Metadata columns - If there are ways you want to organize your data, here is where you add the metadata categories you want to apply to your dataset. 
For each category, type it into the "Add a metadata column" bar and click "Add." Metadata category title must all be lowercase, alphanumeric characters, and cannot have spaces. 
Note: All of the metadata categories will be applied to all of your data (i.e. you cannot just apply the categories to some of your data – you could just put "none" for the ones that don't have a relevant category). 
CSV Template - This is how you are doing to declare the metadata values for all of your documents. 
Click "Create CSV Template" and save the file to a place that is convenient for you. This will create a blank CSV file with all of your documents as rows and metadata categories as columns.
Click "Open CSV Template" and fill out the template with the metadata values of your documents. For ease of filtering, we suggest that you use a consistent set of values across your metadata categories. Example: If you are categorizing your data based on month, it is helpful if multiple files have the "month" metadata value of "june" instead of "June", "june", "JUNE", etc. The entire template must be filled out. If some categories do noy apply to specific files, we suggest filling in "None" or something similar. Be sure to save the file again.
Once the template is filled out, click "Upload CSV Template" and upload your saved and complete CSV file.
Build Qdrant Store - This is where we create the vector database.
Qdrant Folder Path - This is where the vector database will be stored on your computer. Make sure that it is in a location that is easy to find as you will need to access in all future uses. You can change the name as you please. Press 'Build Vector Store'.
Note: As mentioned above, the database will only be on your local computer and not shared with anyone. 
Depending on how large your dataset is, this process might take a while. Do not exit the application. You can run this in the background as you do other things on your computer. Do not close or turn off your computer until the process is complete. You will receive a notification like below indicating the process is complete.

You cannot add new files to an existing vector databse. If you get new data, you will need to rebuild the database from scratch. Therefore, we suggest creating the vector database once with all of your data.

If you are working with collaborators, instead of asking them to re-generate the database, please share the database through secure means such as Stanford Medicine Box or Microsoft Team. 

Steps to share database folder:
Navigate to the folder on your computer.
Compress into a .zip file.
Share on Stanford Medicine box, Microsoft Teams, Google Drive, etc. 

Load DB Option: If already created vector database, please follow these steps:
Select "Load Existing DB Folder"
Press "Select Qdrant Folder" and upload the entire folder that was created during Setup.
Once loaded, a table will appear showing all your .docx documents along with their metadata. Double-check the information to ensure everything looks correct.
""")
        instructions_text.configure(state="disabled")

        # Button to remove/hide the instructions panel
        close_button = ttk.Button(
            right_frame, text="Close Instructions", 
            command=lambda: right_frame.pack_forget()
        )
        close_button.pack(pady=(0,5))

        # ---------------------------------------------------------------------
        # Setup UI on the LEFT frame (unchanged, aside from indentation)
        # ---------------------------------------------------------------------
        mode_frame = ttk.LabelFrame(self.left_frame, text="Load or Create Qdrant DB")
        mode_frame.pack(fill="x", padx=10, pady=5)

        self.db_mode = tk.StringVar(value="load")
        ttk.Radiobutton(mode_frame, text="Load Existing DB Folder", variable=self.db_mode, value="load",
                        command=self.on_db_mode_change).pack(anchor="w", padx=5, pady=2)
        ttk.Radiobutton(mode_frame, text="Create New DB Folder", variable=self.db_mode, value="create",
                        command=self.on_db_mode_change).pack(anchor="w", padx=5, pady=2)

        # Load existing section
        self.load_frame = ttk.Frame(mode_frame)
        self.load_frame.pack(fill="x", padx=5, pady=5)
        ttk.Button(self.load_frame, text="Select Qdrant Folder", command=self.select_qdrant_folder).pack(side="left", padx=5)
        self.load_label = ttk.Label(self.load_frame, text="(No folder selected)")
        self.load_label.pack(side="left", padx=5)

        # Create new section
        self.create_frame = ttk.Frame(mode_frame)
        # --- DOCX Folder Selection ---
        docx_frame = ttk.LabelFrame(self.create_frame, text="Select DOCX Folder")
        docx_frame.pack(fill="x", padx=5, pady=5)
        ttk.Button(docx_frame, text="Select DOCX Folder", command=self.select_docx_folder).pack(anchor="w", padx=5, pady=2)
        self.docx_label = ttk.Label(docx_frame, text="(No folder selected)")
        self.docx_label.pack(anchor="w", padx=5, pady=2)

        # --- Metadata Columns ---
        meta_frame = ttk.LabelFrame(self.create_frame, text="Metadata Columns")
        meta_frame.pack(fill="x", padx=5, pady=5)
        ttk.Label(meta_frame, text="Add a metadata column (e.g., team, role, site_number):").pack(anchor="w", padx=5, pady=2)
        self.meta_entry = ttk.Entry(meta_frame, width=30)
        self.meta_entry.pack(anchor="w", padx=5, pady=2)
        ttk.Button(meta_frame, text="Add", command=self.add_metadata).pack(anchor="w", padx=5, pady=2)
        self.meta_listbox = tk.Listbox(meta_frame, width=40, height=4)
        self.meta_listbox.pack(anchor="w", padx=5, pady=5)

        # --- CSV Template Section ---
        csv_frame = ttk.LabelFrame(self.create_frame, text="CSV Template")
        csv_frame.pack(fill="x", padx=5, pady=5)
        ttk.Button(csv_frame, text="Create CSV Template", command=self.create_csv_template).pack(anchor="w", padx=5, pady=2)
        ttk.Button(csv_frame, text="Open CSV Template", command=self.open_csv_template).pack(anchor="w", padx=5, pady=2)
        ttk.Button(csv_frame, text="Upload CSV", command=self.upload_csv).pack(anchor="w", padx=5, pady=2)
        self.csv_template_label = ttk.Label(csv_frame, text="(No template created)")
        self.csv_template_label.pack(anchor="w", padx=5, pady=2)

        # --- Build Vector Store ---
        build_frame = ttk.LabelFrame(self.create_frame, text="Build Qdrant Store")
        build_frame.pack(fill="x", padx=5, pady=5)
        ttk.Label(build_frame, text="Qdrant Folder Path:").pack(side="left", padx=5)
        self.qdrant_folder_var = tk.StringVar(value=os.path.join(DEFAULT_DOWNLOADS, "my_qdrant_folder"))
        ttk.Entry(build_frame, textvariable=self.qdrant_folder_var, width=40).pack(side="left", padx=5)
        ttk.Button(build_frame, text="Build Vector Store", command=self.build_vector_store).pack(side="left", padx=5)

        # Initially show only the load_frame
        self.create_frame.pack_forget()

    def on_db_mode_change(self):
        mode = self.db_mode.get()
        if mode == "load":
            self.load_frame.pack(fill="x", padx=5, pady=5)
            self.create_frame.pack_forget()
        else:
            self.load_frame.pack_forget()
            self.create_frame.pack(fill="x", padx=5, pady=5)

    def select_qdrant_folder(self):
        folder = filedialog.askdirectory(title="Select Qdrant Folder")
        if not folder:
            return
        json_files = [f for f in os.listdir(folder) if f.lower().endswith(".json")]
        if not json_files:
            messagebox.showerror("Error", "Selected folder does not contain a .json file.")
            return
        coll_path = os.path.join(folder, "collection")
        if not os.path.isdir(coll_path):
            messagebox.showerror("Error", "Selected folder does not contain a 'collection' subfolder.")
            return
        found_sqlite = False
        for sub in os.listdir(coll_path):
            subpath = os.path.join(coll_path, sub)
            if os.path.isdir(subpath):
                if any(f.lower().endswith(".sqlite") for f in os.listdir(subpath)):
                    found_sqlite = True
                    break
        if not found_sqlite:
            messagebox.showerror("Error", "No .sqlite file found in any subfolder of 'collection'.")
            return
        self.controller.qdrant_folder = folder
        self.load_label.config(text=folder)
        try:
            client = QdrantClient(path=folder)
            metadata_keys = set()
            metadata_list = []
            offset = 0
            while True:
                items, next_page = client.scroll(collection_name="qualitative_vector_store", limit=100, offset=offset, with_payload=True, with_vectors=False)
                for item in items:
                    meta = item.payload.get("metadata", {})
                    metadata_keys.update(meta.keys())
                    metadata_list.append(meta)
                if next_page is None:
                    break
                offset = next_page
            self.controller.metadata_types = sorted(list(metadata_keys))
            self.controller.all_doc_metadata = metadata_list
            messagebox.showinfo("DB Loaded", f"Selected Qdrant folder:\n{folder}\nExtracted metadata fields: {', '.join(self.controller.metadata_types)}")
            # Display metadata grid
            self.display_metadata_grid(metadata_list)
            # Refresh Q&A and Analysis Grid drop downs
            if hasattr(self.controller, 'query_page'):
                self.controller.query_page.build_filter_options()
                self.controller.query_page.update_files_tree()
            if hasattr(self.controller, 'analysis_grid_page'):
                self.controller.analysis_grid_page.update_metadata_option_menu()
        except Exception as e:
            messagebox.showerror("Error", f"Failed to extract metadata: {e}")

    def display_metadata_grid(self, metadata_list):
        # First, destroy any old grid frame if it exists
        if hasattr(self, 'metadata_grid_frame') and self.metadata_grid_frame:
            self.metadata_grid_frame.destroy()

        # Now pack the LabelFrame into self.left_frame instead of self
        self.metadata_grid_frame = ttk.LabelFrame(
            self.left_frame, text="Extracted Metadata from Qdrant"
        )
        self.metadata_grid_frame.pack(fill="both", expand=True, padx=10, pady=5)

        # Determine all columns from the metadata
        columns_set = set()
        for meta in metadata_list:
            columns_set.update(meta.keys())
        columns = sorted(list(columns_set))

        # If "filename" is among the columns, move it to the front
        if "filename" in columns:
            columns.remove("filename")
            columns = ["filename"] + columns

        # Collect only unique metadata rows by comparing their column values
        unique_rows = set()
        unique_meta_list = []
        for meta in metadata_list:
            row_tuple = tuple(meta.get(col, "") for col in columns)
            if row_tuple not in unique_rows:
                unique_rows.add(row_tuple)
                unique_meta_list.append(meta)

        # If "filename" is among the columns, sort by filename for consistent display
        if "filename" in columns:
            unique_meta_list = sorted(unique_meta_list, key=lambda x: x.get("filename", ""))

        # Create a Treeview for the metadata
        tree = ttk.Treeview(self.metadata_grid_frame, columns=columns, show="headings")
        for col in columns:
            tree.heading(col, text=col)
            tree.column(col, width=150, anchor="w")

        # Insert only unique rows
        for meta in unique_meta_list:
            row = [meta.get(col, "") for col in columns]
            tree.insert("", tk.END, values=row)

        tree.pack(fill="both", expand=True)

    def select_docx_folder(self):
        folder = filedialog.askdirectory(title="Select Folder with Only DOCX Files")
        if not folder:
            return

        valid_files = []
        for item in os.listdir(folder):
            # Skip hidden or temporary Word files often starting with "~$" or "."
            if item.startswith("~$") or item.startswith("."):
                continue

            full_path = os.path.join(folder, item)

            # Ensure no subfolders
            if os.path.isdir(full_path):
                messagebox.showerror("Error", f"Subfolder found: {item}. Folder must contain only DOCX files.")
                return

            # Ensure file extension is .docx
            if not item.lower().endswith(".docx"):
                messagebox.showerror("Error", f"Non-DOCX file found: {item}")
                return

            # Verify it's a valid docx (zip with 'word/document.xml')
            if not self.is_valid_docx(full_path):
                messagebox.showerror("Error", f"File '{item}' is not a valid DOCX.")
                return

            valid_files.append(item)

        if not valid_files:
            messagebox.showerror("Error", "No valid DOCX files found.")
            return

        self.controller.docx_folder = folder
        self.controller.docx_filenames = sorted([os.path.splitext(fn)[0] for fn in valid_files])
        self.docx_label.config(text=f"{folder} ({len(valid_files)} files)")

    def is_valid_docx(self, filepath: str) -> bool:
        """
        Quick check if a .docx file is valid by confirming it's a zip
        containing 'word/document.xml'.
        """
        try:
            with zipfile.ZipFile(filepath, 'r') as zf:
                return 'word/document.xml' in zf.namelist()
        except Exception:
            return False

    def add_metadata(self):
        # Convert user entry to lowercase immediately
        meta = self.meta_entry.get().strip().lower()

        # Only allow alphanumeric lowercase: letters (a-z) and digits (0-9)
        # allow underscores
        pattern = r'^[a-z0-9_]+$'

        # Check for non-empty input
        if not meta:
            messagebox.showwarning("Input Error", "Please enter a metadata name.")
            return

        # Match to ensure it is strictly letters or digits, all in lowercase
        if not re.match(pattern, meta):
            messagebox.showwarning(
                "Invalid Format",
                "Metadata name must be lowercase letters and digits only, with no spaces or special characters."
            )
            return

        # Prevent duplicates
        if meta in self.controller.metadata_types:
            messagebox.showinfo("Duplicate Metadata", f"'{meta}' is already in the metadata list.")
            return

        # Add valid metadata to the list
        self.controller.metadata_types.append(meta)
        self.meta_listbox.insert(tk.END, meta)
        self.meta_entry.delete(0, tk.END)

    def create_csv_template(self):
        # should save in downloads folder
        downloads_folder = os.path.join(os.path.expanduser("~"), "Downloads")
        if not self.controller.docx_filenames:
            messagebox.showerror("Error", "No DOCX folder selected or folder is empty.")
            return
        if not self.controller.metadata_types:
            messagebox.showwarning("No Metadata", "Please add at least one metadata column first.")
            return
        file_path = filedialog.asksaveasfilename(
            defaultextension=".csv",
            title="Create CSV Template",
            filetypes=[("CSV files", "*.csv")],
            initialdir=downloads_folder
        )
        if not file_path:
            return
        headers = ["filename"] + self.controller.metadata_types
        try:
            with open(file_path, "w", newline="", encoding="utf-8") as f:
                writer = csv.writer(f)
                writer.writerow(headers)
                for fn in self.controller.docx_filenames:
                    row = [fn] + [""] * len(self.controller.metadata_types)
                    writer.writerow(row)
            self.csv_template_label.config(text=f"Template created: {file_path}")
            self.controller.csv_template_path = file_path
        except Exception as e:
            messagebox.showerror("Error", f"Failed to create CSV file: {e}")

    def open_csv_template(self):
        file_path = getattr(self.controller, "csv_template_path", None)
        if not file_path or not os.path.exists(file_path):
            messagebox.showerror("Error", "No CSV template available to open.")
            return
        try:
            import platform
            if os.name == 'nt':
                os.startfile(file_path)
            elif platform.system() == "Darwin":
                subprocess.call(['open', file_path])
            else:
                subprocess.call(['xdg-open', file_path])
        except Exception as e:
            messagebox.showerror("Error", f"Failed to open CSV file: {e}")

    def upload_csv(self):
        file_path = filedialog.askopenfilename(title="Select Completed CSV", filetypes=[("CSV files", "*.csv")])
        if not file_path:
            return
        try:
            df = pd.read_csv(file_path)
        except Exception as e:
            messagebox.showerror("Error", f"Failed to read CSV: {e}")
            return

        # Ensure CSV is not empty
        if df.empty:
            messagebox.showerror("Error", "CSV appears to be empty. Please ensure it has rows and data.")
            return

        # Expected columns: "filename" plus all metadata fields
        expected_cols = ["filename"] + self.controller.metadata_types
        if list(df.columns) != expected_cols:
            messagebox.showerror(
                "Error",
                f"CSV columns must match:\n{expected_cols}\nFound:\n{list(df.columns)}"
            )
            return

        # Make sure every cell is filled out
        for idx, row in df.iterrows():
            for col in expected_cols:
                cell_value = str(row[col]).strip().lower()
                if not cell_value or cell_value in ["nan", "none"]:
                    messagebox.showerror(
                        "Error",
                        f"Column '{col}' has missing or invalid value at row {idx+1}. "
                        "Please fill out all cells."
                    )
                    return

        self.controller.csv_data = df
        messagebox.showinfo("CSV Uploaded", f"CSV uploaded successfully from:\n{file_path}")
        self.csv_template_label.config(text=f"CSV uploaded successfully from:\n{file_path}")

    def build_vector_store(self):
        if not self.controller.docx_filenames:
            messagebox.showerror("Error", "No DOCX folder selected.")
            return
        if self.controller.csv_data is None or self.controller.csv_data.empty:
            messagebox.showerror("Error", "CSV file not validated or is empty. Make sure you've uploaded and validated it.")
            return
        folder = self.qdrant_folder_var.get().strip()
        if not folder:
            messagebox.showwarning("Input Error", "Please specify a Qdrant folder path.")
            return
        
        # Create folder if it doesn't exist
        try:
            if not os.path.exists(folder):
                os.makedirs(folder, exist_ok=True)
        except Exception as e:
            messagebox.showerror("Error", f"Failed to create folder: {e}")
            return
        
        print("[VectorStore] Starting build process...")
        all_texts = []
        all_metadatas = []
        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)

        # Map from filename -> metadata
        csv_map = {}
        for _, row in self.controller.csv_data.iterrows():
            fname = str(row["filename"]).strip()
            meta = {}
            for c in self.controller.metadata_types:
                value = row[c]
                if pd.isna(value):
                    value = ""
                meta[c] = str(value).strip()
            csv_map[fname] = meta
        
        # Read and split DOCX files
        for file in os.listdir(self.controller.docx_folder):
            if file.startswith("~$") or file.startswith("."):
                continue
            if file.lower().endswith(".docx"):
                base = os.path.splitext(file)[0]
                if base not in csv_map:
                    print(f"[VectorStore] Warning: {file} (base: {base}) not found in CSV. Skipping.")
                    continue
                full_path = os.path.join(self.controller.docx_folder, file)
                print(f"[VectorStore] Processing file: {file}")
                content = read_docx_file(full_path)
                chunks = splitter.split_text(content)
                for idx, ch in enumerate(chunks, start=1):
                    all_texts.append(ch)
                    md = {"filename": base}
                    md.update(csv_map[base])
                    all_metadatas.append(md)
        
        if not all_texts:
            messagebox.showerror("Error", "No text found from DOCX files.")
            return
        
        print(f"[VectorStore] Creating Qdrant collection in folder: {folder}")
        try:
            client = QdrantClient(path=folder)
            try:
                client.delete_collection(collection_name="qualitative_vector_store")
                print("[VectorStore] Deleted old collection.")
            except Exception as e:
                print("[VectorStore] No old collection to delete or error during deletion:", e)
            
            client.create_collection(
                collection_name="qualitative_vector_store",
                vectors_config=VectorParams(size=3072, distance=Distance.COSINE)
            )
            print("[VectorStore] Collection created.")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to create Qdrant collection: {e}")
            return
        
        # Create embeddings
        embeddings = StanfordSecureEmbeddings()

        # Instantiate QdrantVectorStore
        vectorstore = QdrantVectorStore(
            client=client,
            collection_name="qualitative_vector_store",
            embeddings=embeddings,
        )

        # Add texts in batches
        total_chunks = len(all_texts)
        batch_size = 50
        num_batches = ((total_chunks - 1) // batch_size) + 1
        print(f"[VectorStore] Adding {total_chunks} text chunks in batches of {batch_size}...")
        for i in range(0, total_chunks, batch_size):
            batch_texts = all_texts[i : i + batch_size]
            batch_metadata = all_metadatas[i : i + batch_size]
            batch_index = i // batch_size + 1
            print(f"[VectorStore] Adding batch {batch_index} of {num_batches} ...")
            try:
                vectorstore.add_texts(texts=batch_texts, metadatas=batch_metadata)
            except Exception as e:
                messagebox.showerror("Error", f"Error adding batch {batch_index}: {e}")
                return
            if batch_index % 3 == 0:
                try:
                    items, _ = client.scroll(collection_name="qualitative_vector_store", limit=3, with_payload=True, with_vectors=False)
                    print("[VectorStore] Sample collection data:")
                    # for item in items:
                    #     print(item.payload)
                except Exception as ex:
                    print("[VectorStore] Error fetching sample data:", ex)

        print("[VectorStore] All text chunks added successfully.")
        
        # --------------------------------------------------------
        # AUTOMATICALLY LOAD THE NEWLY CREATED DB
        # (Equivalent to select_qdrant_folder logic, minus file dialog)
        # --------------------------------------------------------
        try:
            # Save references so other pages can use them
            self.controller.qdrant_folder = folder
            self.controller.vector_store = vectorstore
            
            # Extract metadata
            metadata_keys = set()
            metadata_list = []
            offset = 0
            while True:
                items, next_page = client.scroll(
                    collection_name="qualitative_vector_store", 
                    limit=100, 
                    offset=offset, 
                    with_payload=True, 
                    with_vectors=False
                )
                for item in items:
                    meta = item.payload.get("metadata", {})
                    metadata_keys.update(meta.keys())
                    metadata_list.append(meta)
                if next_page is None:
                    break
                offset = next_page
            
            self.controller.metadata_types = sorted(list(metadata_keys))
            self.controller.all_doc_metadata = metadata_list
            
            # Display metadata in the UI (like in select_qdrant_folder)
            self.display_metadata_grid(metadata_list)

            # Notify other tabs to refresh
            if hasattr(self.controller, 'query_page'):
                self.controller.query_page.build_filter_options()
                self.controller.query_page.update_files_tree()
            if hasattr(self.controller, 'analysis_grid_page'):
                self.controller.analysis_grid_page.update_metadata_option_menu()

            messagebox.showinfo(
                "Success", 
                f"Qdrant store created and loaded from folder:\n{folder}\n\n"
                f"Extracted metadata fields: {', '.join(self.controller.metadata_types)}"
            )
        except Exception as e:
            messagebox.showerror("Error", f"Failed to load newly created DB: {e}")

######################################################
# PAGE 3: QUERY TAB (formerly Q&A)
######################################################
class QueryTab(ttk.Frame):
    def __init__(self, parent, controller: App):
        super().__init__(parent)
        self.controller = controller


        # ---------------------------------------------------------------------
        # Outer container that splits the page into a left Frame (main UI)
        # and a right Frame (instructions sidebar).
        # ---------------------------------------------------------------------
        outer_frame = ttk.Frame(self)
        outer_frame.pack(fill="both", expand=True)

        # # Left side for your existing Query UI
        # left_frame = ttk.Frame(outer_frame)
        # left_frame.pack(side="left", fill="both", expand=True, padx=5, pady=5)
        self.left_scrollable = ScrollableFrame(outer_frame)
        self.left_scrollable.pack(side="left", fill="both", expand=True)
        
        # Access the scrollable frame for adding your controls
        container = self.left_scrollable.scrollable_frame


        # Right side for the scrollable instructions (omitted here for brevity).
        # You can place your instructions or any other widgets on the right side.
        right_frame = ttk.Frame(outer_frame, width=400)
        right_frame.pack(side="right", fill="y", padx=5, pady=5)

        
         # ---------------------------------------------------------------------
        # Instructions for the Query Tab (scrollable, closeable)
        # ---------------------------------------------------------------------
        instructions_label = ttk.Label(
            right_frame, text="Query Tab Instructions",
            font=("Helvetica", 14, "bold")
        )
        instructions_label.pack(pady=5)

        instructions_text = scrolledtext.ScrolledText(right_frame, width=40, wrap="word")
        instructions_text.config(font=("Helvetica", 15))
        instructions_text.pack(fill="both", expand=True, padx=5, pady=(0,5))

        # Insert the instructions exactly as provided - DO NOT change wording:
        instructions_text.insert("1.0", """This screen functions similarly to an LLM web interface like ChatGPT. The key features are: 
Retrieval-Augmented Generation (RAG) - Your queries are grounded on the data most relevant to the question, improving accuracy and relevance.
Metadata filtering - Easily narrow your search to specific documents using the metadata you defined during setup.
Fine-Tuned Model Controls - Advanced settings allow deeper control over the model's behavior (details on this below).

Filter Options 
Use the filters to select which subset of your documents you want the model to search from. These filters are based on your custom metadata fields (e.g., month, site, interviewer).
The "Selected Files" section displats which files will be used for the query.

Excerpt Retrieval Settings
In Retrieval Augmented Generation (RAG), the model retrieves the most relevant text excerpts from your dataset and provides them alongside your question to generate a high-quality response.

Score Threshold 
The Score Threshold determines how closely an excerpt must match your question to be included in the query. It ranges from 0.0 to 1.0: 1.0 represents a perfect match and 0.0 represents no match or similarity. Empirically, we found a threshold of 0.4 tends to capture excerpts that are meaningfully related to the query. The app defaults to 0.3, meaning excerpts with a similarity score between 0.3 and 1.0 will be included, sorted by relevance (highest first).

Number of Results
You can also choose how many relevant excerpts you want to provide to the LLM with the "Number of Results." This limits the number of results to your chosen number if there are more excerpts than the number you set that are within your score threshold. Excerpts are always ranked and selected by highest similarity score.
Example: Imagine you select a score threshold of 0.5 and number of results to 10. If there are 20 excerpts with a score greater than 0.5, then only the first 10 that have the highest score will be chosen for the query to the LLM. If there are 5 excerpts with a score greater than 0.5, then only the five will be chosen for the query to the LLM.

After submitting a question, you'll be able to view the retrieved excerpts at the bottom of the screen. This lets you verify exactly what context was provided to the LLM.

LLM Settings 
These settings allow you to choose which large language model (LLM) to use and how you want it to behave when generating answers.

Model Selection
We currently support OpenAI's ChatGPT-4o, OpenAI's o1-mini, OpenAI's o1, and deepseek R1 LLM models (good comparison here). In a nutshell, the order of models in terms of performance (best to worst) is OpenAI's o1, deepseek R1, ChatGPT-4o, OpenAI's o1-mini. ChatGPT-4o is a next-word generator which means that it develops its answer by picking the next best word based on all of its training. OpenAI's o1-mini, OpenAI's o1, and deepseek R1 have reasoning capability which means they take longer to "think" about the question and answer. OpenAI's o1-mini is trained on a smaller amount of data than OpenAI's o1. OpenAI's o1 and deepseek R1 are competing models from different companies with similar performance. deepseek R1 costs less money per query than OpenAI's o1 model, and ChatGPT-4o is the cheapest (i.e. best for many queries). We encourage you to play around with which model works best for your needs.

Temperature 
The temperature setting controls how random or creative the model's responses are. Temperature is on a scale of 0.0-1.0. A temperature of 0.0 is more deterministic and factual. The model will pick the most likely next word every time. This is useful when accuracy and consistency are top priorities (e.g., summaries, technical info). A temperature of 0.5 is a balance of logic and creativity. This is best for conversational or open-ended answers. A temperature of 1.0 is more creative and diverse. The model becomes more random and imaginative. This is great for brainstorming, poetry, or creative writing.

ChatGPT is typically set at 0.7 and deepseek is typically set at 0.5-0.6. We suggest lower values to reduce hallucination. "Hallucination" refers to when the model makes up information that sounds correct but isn't true. Lower temperatures reduce this risk. More information here. The default setting is 0.0.

Max Tokens 
The Max Tokens setting controls how long the LLM's response can be. 

A token is a small unit of text that the language model uses to read and generate language. It could be a whole word (like 'hello'), part of a word (like 'walk' and 'ing'), or even punctuation (like '.' or ','). On average, 1 token is about 4 characters or ¾ of a word in English. Token limits determine how much you can input (your prompt), how long the model's response can be, and how much the query costs.

This setting determines how long the model's response can be and how much it will cost (more tokens = longer responses = higher cost). If you increase the token count, you will allow for longer responses. If your prompt and expected response exceed the model's token limit, the model may cut off part of the answer or skip important context from your prompt. The default setting for max tokens is 2,000. Stanford SecureGPT has a limit of 4,000 maximum tokens.

Prompt Instructions & Response Format
This section defines exactly what is sent to the LLM when you ask a question. The model receives a combination of prompt instructions, response format, question, and retrieved excerpts.

Prompt Instructions 
This is where you guide the LLM on how to respond.
You can include: the type of response (e.g., summary, bullet points, 3-sentence answer), broader context about the study or purpose of analysis, or one or structure preferences (e.g., concise, formal). Default instructions are provided but can be customized for each question. 

Response Format
                                 
This sets the structure of the answer you want.

The more specific your format, the easier it is to get a consistent and useful response. The LLM typically adheres to this format, but may not exactly. Examples: a list of key themes, paragraph followed by bullet points, summary point followed by quote. 

We provide a default format you can change.
Question
This is the actual query to the LLM based on your dataset.

Ask questions about content within the documents, such as: "How do participants describe their diabetes journey?" or "What barriers to healthcare access are mentioned?" Do not ask about file statistics or metadata-level info like: "How many interviews are with women?" or "Which site has the most documents?".

The app first retrieves relevant text excerpts from your selected documents based on your question. Then it packages everything (prompt instructions, response format, question, and excerpts) and sends it to the LLM.

Once you've entered your prompt instructions, response format, and question, press "Submit Question" at the bottom.

After submitting the question, you will see the retrieved excerpts and the LLM response based on the retrieved excerpts. Note: If you use deepseek R1, you can see how the LLM "thinks" with the information in enclosed by <think> and </think>.

If you want to save the excerpts and LLM response, click "Save to CSV" and choose a csv file to save the response to. It will create a new row in the CSV file with the question, excerpts, and LLM response.  Questions and settings are not saved automatically. They will be lost if you ask a new question or close the app.
""")
        
        instructions_text.configure(state="disabled")

        close_button = ttk.Button(right_frame, text="Close Instructions",
                                  command=lambda: right_frame.pack_forget())
        close_button.pack(pady=(0,5))

        # mode_frame = ttk.LabelFrame(left_frame, text="Load or Create Qdrant DB")
        # mode_frame.pack(fill="x", padx=10, pady=5)

        # self.scrollable = ScrollableFrame(self)
        # # self.scrollable.pack(fill="x", padx=10, pady=5)
        # self.scrollable = ScrollableFrame(left_frame)
        # self.scrollable.pack(fill="both", expand=True)
        # container = self.scrollable.scrollable_frame


        # Filter Options as dropdowns
        # ---------------------------------------------------------------------
        self.filter_frame = ttk.LabelFrame(container, text="Filter Options")
        self.filter_frame.pack(fill="x", padx=10, pady=5)
        self.filter_comboboxes = {}
        self.build_filter_options()

        # ---------------------------------------------------------------------
        # Files Treeview renamed to "Selected Files"
        # ---------------------------------------------------------------------
        self.files_frame = ttk.LabelFrame(container, text="Selected Files")
        self.files_frame.pack(fill="both", expand=True, padx=10, pady=5)

        self.files_tree = ttk.Treeview(self.files_frame, columns=[], show="headings")
        self.files_tree.pack(fill="both", expand=True)

        # files_scrollbar = ttk.Scrollbar(self.files_frame, orient="vertical", command=self.files_tree.yview)
        # files_scrollbar.pack(side="right", fill="y")
        # self.files_tree.configure(yscrollcommand=files_scrollbar.set)

        files_scrollbar = ttk.Scrollbar(self.files_frame, orient="vertical", command=self.files_tree.yview)
        files_scrollbar.pack(side="right", fill="y")
        files_scrollbar_x = ttk.Scrollbar(self.files_frame, orient="horizontal", command=self.files_tree.xview)
        files_scrollbar_x.pack(side="bottom", fill="x")

        self.files_tree.configure(
            yscrollcommand=files_scrollbar.set,
            xscrollcommand=files_scrollbar_x.set
        )

        # ---------------------------------------------------------------------
        # Excerpt Retrieval Settings
        # ---------------------------------------------------------------------
        ret_frame = ttk.LabelFrame(container, text="Excerpt Retrieval Settings")
        ret_frame.pack(fill="x", padx=10, pady=5)

        ttk.Label(ret_frame, text="Score Threshold (0.0 - 1.0):")\
            .grid(row=0, column=0, padx=5, pady=5, sticky="e")
        self.score_thresh_entry = ttk.Entry(ret_frame)
        self.score_thresh_entry.insert(0, "0.3")
        self.score_thresh_entry.grid(row=0, column=1, padx=5, pady=5)

        ttk.Label(ret_frame, text="Number of Results:")\
            .grid(row=1, column=0, padx=5, pady=5, sticky="e")
        self.num_results_entry = ttk.Entry(ret_frame)
        self.num_results_entry.insert(0, "20")
        self.num_results_entry.grid(row=1, column=1, padx=5, pady=5)

        # ---------------------------------------------------------------------
        # LLM Settings
        # ---------------------------------------------------------------------
        llm_frame = ttk.LabelFrame(container, text="LLM Settings")
        llm_frame.pack(fill="x", padx=10, pady=5)

        ttk.Label(llm_frame, text="Model Selection:")\
            .grid(row=0, column=0, padx=5, pady=5, sticky="e")

        self.model_var = tk.StringVar(value="gpt4o")
        model_opts = [("gpt-4o", "gpt4o"), ("o1-mini", "o1-mini"), ("o1", "o1"), ("DeepSeek R1", "deepseek")]
        col = 1
        for lbl, val in model_opts:
            ttk.Radiobutton(llm_frame, text=lbl, variable=self.model_var, value=val)\
                .grid(row=0, column=col, padx=5, pady=5)
            col += 1

        ttk.Label(llm_frame, text="Temperature (0.0 - 1.0):")\
            .grid(row=1, column=0, padx=5, pady=5, sticky="e")
        self.temp_entry = ttk.Entry(llm_frame)
        self.temp_entry.insert(0, "0.0")
        self.temp_entry.grid(row=1, column=1, padx=5, pady=5)

        ttk.Label(llm_frame, text="Max Tokens:")\
            .grid(row=1, column=2, padx=5, pady=5, sticky="e")
        self.max_tokens_entry = ttk.Entry(llm_frame)
        self.max_tokens_entry.insert(0, "2000")
        self.max_tokens_entry.grid(row=1, column=3, padx=5, pady=5)

        # ---------------------------------------------------------------------
        # Prompt Instructions & Response Format
        # ---------------------------------------------------------------------
        prompt_frame = ttk.LabelFrame(container, text="Prompt Instructions & Response Format")
        prompt_frame.pack(fill="x", padx=10, pady=5)

        ttk.Label(prompt_frame, text="Instructions:").pack(anchor="w", padx=5, pady=2)
        self.instructions_text = scrolledtext.ScrolledText(prompt_frame, height=4)
        default_instructions = (
            "Provide a concise and specific answer in 3-5 bullet points. Each bullet point must include:\n"
            "1. A clear summary of the insight or opinion in a single sentence.\n"
            "2. A full supporting quote from the data, written verbatim, with no omissions or ellipses. "
            "Ensure the quote is at least 2-3 sentences long and provides enough context for the reader to "
            "understand its meaning fully.\n"
            "3. Ensure the entire output adheres to the above criteria."
        )
        self.instructions_text.insert("1.0", default_instructions)
        self.instructions_text.pack(fill="x", padx=5, pady=2)

        ttk.Label(prompt_frame, text="Response Format:").pack(anchor="w", padx=5, pady=2)
        self.response_format_text = scrolledtext.ScrolledText(prompt_frame, height=3)
        default_format = "[Summary Point]: Description of strategy or structure.\n[Quote]: From [All Metadata]"
        self.response_format_text.insert("1.0", default_format)
        self.response_format_text.pack(fill="x", padx=5, pady=2)

        # ---------------------------------------------------------------------
        # Question Panel
        # ---------------------------------------------------------------------
        q_frame = ttk.LabelFrame(container, text="Question")
        q_frame.pack(fill="x", padx=10, pady=5)

        self.question_text = scrolledtext.ScrolledText(q_frame, height=4)
        self.question_text.pack(fill="x", padx=5, pady=5)

        # ---------------------------------------------------------------------
        # Response Panels
        # ---------------------------------------------------------------------
        display_frame = ttk.Frame(container)
        display_frame.pack(fill="both", expand=True, padx=10, pady=5)

        excerpts_frame = ttk.LabelFrame(display_frame, text="Retrieved Excerpts (Context)")
        excerpts_frame.pack(side="left", fill="both", expand=True, padx=5, pady=5)
        self.excerpts_text = scrolledtext.ScrolledText(excerpts_frame, height=10, width=50)
        self.excerpts_text.pack(fill="both", expand=True, padx=5, pady=5)

        answer_frame = ttk.LabelFrame(display_frame, text="LLM Answer")
        answer_frame.pack(side="left", fill="both", expand=True, padx=5, pady=5)
        self.answer_text = scrolledtext.ScrolledText(answer_frame, height=10, width=50)
        self.answer_text.pack(fill="both", expand=True, padx=5, pady=5)

        # ---------------------------------------------------------------------
        # Buttons: Submit Query and Save Query to CSV
        # ---------------------------------------------------------------------
        button_frame = ttk.Frame(container)
        button_frame.pack(pady=5)

        submit_button = ttk.Button(button_frame, text="Submit Question", command=self.submit_question)
        submit_button.pack(side="left", padx=5)

        save_button = ttk.Button(button_frame, text="Save Query to CSV", command=self.save_results_to_csv)
        save_button.pack(side="left", padx=5)

    def build_filter_options(self):
        for widget in self.filter_frame.winfo_children():
            widget.destroy()
        self.filter_comboboxes = {}
        metadata_fields = sorted({key for meta in self.controller.all_doc_metadata for key in meta.keys()})
        for i, field in enumerate(metadata_fields):
            label = ttk.Label(self.filter_frame, text=field)
            label.grid(row=0, column=i, padx=5, pady=5)
            values = sorted({meta.get(field, "") for meta in self.controller.all_doc_metadata if meta.get(field, "")})
            combobox = ttk.Combobox(self.filter_frame, values=["All"] + values, state="readonly")
            combobox.current(0)
            combobox.grid(row=1, column=i, padx=5, pady=5)
            combobox.bind("<<ComboboxSelected>>", self.update_filters)
            self.filter_comboboxes[field] = combobox

    def update_filters(self, event=None):
        current_filters = {field: combobox.get() for field, combobox in self.filter_comboboxes.items()}
        filtered = self.controller.all_doc_metadata
        for field, value in current_filters.items():
            if value != "All":
                filtered = [meta for meta in filtered if meta.get(field, "") == value]
        for field, combobox in self.filter_comboboxes.items():
            values = sorted({meta.get(field, "") for meta in filtered if meta.get(field, "")})
            current = combobox.get()
            combobox['values'] = ["All"] + values
            if current not in ["All"] + values:
                combobox.set("All")
        self.update_files_tree(filtered)

    def update_files_tree(self, metadata_list=None):
        for item in self.files_tree.get_children():
            self.files_tree.delete(item)
        if metadata_list is None:
            metadata_list = self.controller.all_doc_metadata
        if not metadata_list:
            return
        columns = sorted({key for meta in metadata_list for key in meta.keys()})
        if "filename" in columns:
            columns.remove("filename")
            columns = ["filename"] + columns
            metadata_list = sorted(metadata_list, key=lambda x: x.get("filename", ""))
        self.files_tree["columns"] = columns
        for col in columns:
            self.files_tree.heading(col, text=col)
            self.files_tree.column(col, width=150, anchor="w")
        seen_rows = set()
        for meta in metadata_list:
            row_tuple = tuple(meta.get(col, "") for col in columns)
            if row_tuple not in seen_rows:
                seen_rows.add(row_tuple)
                # Insert the row into the Treeview
                self.files_tree.insert("", "end", values=row_tuple)

    def save_results_to_csv(self):
        """
        Allows user to pick or create a CSV file in their Downloads folder.
        If the file already exists, user can choose to append or overwrite.
        Saves the query results to the chosen CSV.
        """
        # Default the file dialog to the Downloads folder
        downloads_folder = os.path.join(os.path.expanduser("~"), "Downloads")

        # Ask for a filename; if the user picks an existing CSV or types a new name,
        # we'll handle that below
        file_path = filedialog.asksaveasfilename(
            title="Save Query Results to CSV",
            defaultextension=".csv",
            filetypes=[("CSV Files", "*.csv")],
            initialdir=downloads_folder,
            initialfile="my_query_results.csv"  # optional default name
        )

        if not file_path:
            # User canceled
            return

        # Check if the file exists; if so prompt user to append or overwrite
        mode = 'w'  # default is write (creating new or overwriting)
        write_header = True
        if os.path.exists(file_path):
            # Ask if user wants to append to the existing file
            choice = messagebox.askyesno(
                "File Exists",
                ("The file you selected already exists.\n\n"
                 "Click 'Yes' to append to the existing file.\n"
                 "Click 'No' to overwrite it.")
            )
            if choice:
                mode = 'a'   # append
                write_header = False  # typically skip header if appending
            else:
                mode = 'w'   # overwrite
                write_header = True

        # Here, you define the column headers for your query results
        headers = ["filename", "some_metadata", "another_field", "excerpt", "response"]

        try:
            with open(file_path, mode=mode, newline='', encoding="utf-8") as f:
                writer = csv.writer(f)
                # Write header only if mode is 'w' or user chose to overwrite
                if write_header:
                    writer.writerow(headers)

                # Example data: adapt to your own query results
                # For instance, you might have self.current_query_results as a list of dicts
                sample_rows = [
                    ["doc1", "meta1", "meta2", "Excerpt from doc1...", "LLM Response snippet..."],
                    ["doc2", "metaA", "metaB", "Excerpt from doc2...", "LLM Response..."]
                ]

                # Write the rows
                for row in sample_rows:
                    writer.writerow(row)

            messagebox.showinfo("Success", f"Query results saved to:\n{file_path}")

        except Exception as e:
            messagebox.showerror("Error", f"Failed to save CSV:\n{e}")

    def submit_question(self):
        if not self.controller.qdrant_folder:
            messagebox.showwarning("No DB", "Please load or create a Qdrant DB first.")
            return
        if not self.controller.vector_store:
            try:
                client = QdrantClient(path=self.controller.qdrant_folder)
                self.controller.vector_store = QdrantVectorStore(
                    client=client,
                    collection_name="qualitative_vector_store",
                    embeddings=StanfordSecureEmbeddings()
                )
            except Exception as e:
                messagebox.showerror("Error", f"Failed to connect to Qdrant: {e}")
                return
        question = self.question_text.get("1.0", tk.END).strip()
        if not question:
            messagebox.showwarning("Input Error", "Please enter a question.")
            return
        filters = []
        for field, combobox in self.filter_comboboxes.items():
            value = combobox.get()
            if value != "All":
                filters.append(FieldCondition(key=f"metadata.{field}", match=MatchValue(value=value)))
        q_filter = Filter(must=filters) if filters else None
        print(q_filter)
        try:
            query_vec = StanfordSecureEmbeddings().embed_query(question)
        except Exception as e:
            messagebox.showerror("Error", f"Embedding error: {e}")
            return
        try:
            max_tokens = int(self.max_tokens_entry.get().strip())
            if max_tokens > 4000:
                messagebox.showerror("Error", "Maximum tokens cannot exceed 4000.")
                return
        except ValueError:
            messagebox.showerror("Error", "Invalid number for max tokens.")
            return
        temperature = float(self.temp_entry.get().strip())
        model_choice = self.model_var.get()
        try:
            limit_ = int(self.num_results_entry.get())
            thresh_ = float(self.score_thresh_entry.get())
            results = self.controller.vector_store.client.search(
                collection_name="qualitative_vector_store",
                query_vector=query_vec,
                limit=limit_,
                with_payload=True,
                query_filter=q_filter,
                score_threshold=thresh_
            )
            # print("collection_name: ", "qualitative_vector_store")
            # print("query_vector: ", query_vec)
            # print("results_limit: ", limit_)
            # print("with_payload: ", True)
            # print("query_filter: ", q_filter)
            # print("score_threshold: ", thresh_)

            #print(limit_, thresh_, results)
        except Exception as e:
            messagebox.showerror("Error", f"Search failed: {e}")
            return
        self.excerpts_text.delete("1.0", tk.END)
        self.answer_text.delete("1.0", tk.END)
        if not results:
            self.answer_text.insert(tk.END, "No relevant documents found.")
            return
        context_str = ""
        for r in results:
            meta = r.payload.get("metadata", {})
            chunk = r.payload.get("page_content", "")
            metadata_items = [f"{k}: {v}" for k, v in meta.items()]
            metadata_str = "; ".join(metadata_items)
            context_str += f"[Score: {r.score:.4f}, {metadata_str}]\n{chunk}\n\n"
        self.excerpts_text.insert(tk.END, context_str)
        instructions = self.instructions_text.get("1.0", tk.END).strip()
        resp_format = self.response_format_text.get("1.0", tk.END).strip()
        final_prompt = (
            f"Instructions:\n{instructions}\n\n"
            f"Response Format:\n{resp_format}\n\n"
            f"Context:\n{context_str}\n\n"
            f"Question: {question}\nAnswer:"
        )
        if model_choice in ["gpt4o", "o1-mini", "o1"]:
            llm = StanfordSecureGPT(temperature=temperature, max_tokens=max_tokens, model_selection=model_choice)
        elif model_choice == "deepseek":
            llm = DeepSeekR1Client(temperature=temperature, max_tokens=max_tokens)
        else:
            self.answer_text.insert(tk.END, "Unsupported model.")
            return
        try:
            answer = llm.invoke(final_prompt)
            self.answer_text.insert(tk.END, answer)
        except Exception as e:
            self.answer_text.insert(tk.END, f"LLM call failed: {e}")

######################################################
# PAGE 4: ANALYSIS GRID TAB
######################################################
class AnalysisGridTab(ttk.Frame):
    def __init__(self, parent, controller: App):
        super().__init__(parent)
        self.controller = controller

          # ---------------------------------------------------------------------
        # Outer container that splits the page into a left Frame (main UI)
        # and a right Frame (instructions sidebar).
        # ---------------------------------------------------------------------
        outer_frame = ttk.Frame(self)
        outer_frame.pack(fill="both", expand=True)


        self.left_scrollable = ScrollableFrame(outer_frame)
        self.left_scrollable.pack(side="left", fill="both", expand=True)

        # Access the scrollable frame for adding your controls
        container = self.left_scrollable.scrollable_frame



        # # Left side for your existing Query UI
        # left_frame = ttk.Frame(outer_frame)
        # left_frame.pack(side="left", fill="both", expand=True, padx=5, pady=5)

        # Right side for the scrollable instructions (omitted here for brevity).
        # You can place your instructions or any other widgets on the right side.
        right_frame = ttk.Frame(outer_frame, width=400)
        right_frame.pack(side="right", fill="y", padx=5, pady=5)

        
         # ---------------------------------------------------------------------
        # Instructions for the Query Tab (scrollable, closeable)
        # ---------------------------------------------------------------------
        instructions_label = ttk.Label(
            right_frame, text="Analysis Grid Instructions",
            font=("Helvetica", 14, "bold")
        )
        instructions_label.pack(pady=5)

        instructions_text = scrolledtext.ScrolledText(right_frame, width=40, wrap="word")
        instructions_text.config(font=("Helvetica", 15))
        instructions_text.pack(fill="both", expand=True, padx=5, pady=(0,5))

        # Insert the instructions exactly as provided - DO NOT change wording:
        instructions_text.insert("1.0", """
This is helpful for asking a set of questions across all of your data. You can partition your data based on the metadata categories that you provided and then ask multiple questions across those categories. Each row is the set of data that represent one metadata value in your metadata category and each column is a question. You can export this grid to a CSV file.

Excerpt Retrieval Settings
In Retrieval Augmented Generation (RAG), the model retrieves the most relevant text excerpts from your dataset and provides them alongside your question to generate a high-quality response.
Score Threshold 
The Score Threshold determines how closely an excerpt must match your question to be included in the query. It ranges from 0.0 to 1.0: 1.0 represents a perfect match and 0.0 represents no match or similarity. Empirically, we found a threshold of 0.4 tends to capture excerpts that are meaningfully related to the query. The app defaults to 0.3, meaning excerpts with a similarity score between 0.3 and 1.0 will be included, sorted by relevance (highest first).
Number of Results
You can also choose how many relevant excerpts you want to provide to the LLM with the "Number of Results." This limits the number of results to your chosen number if there are more excerpts than the number you set that are within your score threshold. Excerpts are always ranked and selected by highest similarity score. 

Example: Imagine you select a score threshold of 0.5 and number of results to 10. If there are 20 excerpts with a score greater than 0.5, then only the first 10 that have the highest score will be chosen for the query to the LLM. If there are 5 excerpts with a score greater than 0.5, then only the five will be chosen for the query to the LLM. 

After submitting a question, you'll be able to view the retrieved excerpts at the bottom of the screen. This lets you verify exactly what context was provided to the LLM.
LLM Settings 
These settings allow you to choose which large language model (LLM) to use and how you want it to behave when generating answers.
Model Selection
We currently support OpenAI's ChatGPT-4o, OpenAI's o1-mini, OpenAI's o1, and deepseek R1 LLM models (good comparison here). In a nutshell, the order of models in terms of performance (best to worst) is OpenAI's o1, deepseek R1, ChatGPT-4o, OpenAI's o1-mini. ChatGPT-4o is a next-word generator which means that it develops its answer by picking the next best word based on all of its training. OpenAI's o1-mini, OpenAI's o1, and deepseek R1 have reasoning capability which means they take longer to "think" about the question and answer. OpenAI's o1-mini is trained on a smaller amount of data than OpenAI's o1. OpenAI's o1 and deepseek R1 are competing models from different companies with similar performance. deepseek R1 costs less money per query than OpenAI's o1 model, and ChatGPT-4o is the cheapest (i.e. best for many queries). We encourage you to play around with which model works best for your needs.
Temperature 
The temperature setting controls how random or creative the model's responses are. Temperature is on a scale of 0.0-1.0. A temperature of 0.0 is more deterministic and factual. The model will pick the most likely next word every time. This is useful when accuracy and consistency are top priorities (e.g., summaries, technical info). A temperature of 0.5 is a balance of logic and creativity. This is best for conversational or open-ended answers. A temperature of 1.0 is more creative and diverse. The model becomes more random and imaginative. This is great for brainstorming, poetry, or creative writing.


ChatGPT is typically set at 0.7 and deepseek is typically set at 0.5-0.6. We suggest lower values to reduce hallucination. "Hallucination" refers to when the model makes up information that sounds correct but isn't true. Lower temperatures reduce this risk. More information here. The default setting is 0.0. 
Max Tokens 
The Max Tokens setting controls how long the LLM's response can be. 

A token is a small unit of text that the language model uses to read and generate language. It could be a whole word (like 'hello'), part of a word (like 'walk' and 'ing'), or even punctuation (like '.' or ','). On average, 1 token is about 4 characters or ¾ of a word in English. Token limits determine how much you can input (your prompt), how long the model's response can be, and how much the query costs.

This setting determines how long the model's response can be and how much it will cost (more tokens = longer responses = higher cost). If you increase the token count, you will allow for longer responses. If your prompt and expected response exceed the model's token limit, the model may cut off part of the answer or skip important context from your prompt. The default setting for max tokens is 2,000. Stanford SecureGPT has a limit of 4,000 maximum tokens. 
Prompt Instructions & Response Format
This section defines exactly what is sent to the LLM when you ask a question. The model receives a combination of prompt instructions, response format, question, and retrieved excerpts. 
Prompt Instructions 
This is where you guide the LLM on how to respond.
You can include: the type of response (e.g., summary, bullet points, 3-sentence answer), broader context about the study or purpose of analysis, or one or structure preferences (e.g., concise, formal). Default instructions are provided but can be customized for each question. 

Response Format
This sets the structure of the answer you want.

The more specific your format, the easier it is to get a consistent and useful response. The LLM typically adheres to this format, but may not exactly. Examples: a list of key themes, paragraph followed by bullet points, summary point followed by quote. 


We provide a default format you can change.
Data Segmentation 
This feature allows you to break down your dataset by metadata categories you defined earlier (e.g., gender, clinic, month) and analyze each group separately.

It will load rows which are the relevant values for the metadata categories. For example, if you want to categorize your data based on "gender," it will load "female," "male," and "nonbinary" as the rows. 

Click "Load Segment" to load into the analysis grid below.
Add Question
This is where you type in your question to the LLM. Do not include instructions here. For example, you may ask "How do the sources feel about their diabetes condition?" or "What does this source say about healthcare management challenges?". You cannot ask questions generally about the statistics or structure of your dataset such as "How many interviews do I have with female participants?". With your question, we will first retrieve the relevant text data from the filtered files based on the question. Then, we give the prompt instructions, response format, question, and retrieved data to the LLM.

After each question, press "Add question" to load into the grid. This adds your question to the analysis grid. You can repeat this to add multiple questions.

Once you have added all of the questions, click "Fill in Grid." This will run each question across each metadata value group. The app sends the prompt instructions, response format, question, and excerpts to the LLM for every cell in the grid.

Depending on the number of questions, groups, and model used, this process may take time.
Do not close the app or your computer during generation.

After the grid is populated, click "Save to CSV". Choose or create a file to save the results. The analysis grid and your settings are not saved automatically. Be sure to export your work before exiting the application.

""")
        
        instructions_text.configure(state="disabled")

        close_button = ttk.Button(right_frame, text="Close Instructions",
                                  command=lambda: right_frame.pack_forget())
        close_button.pack(pady=(0,5))
        
        # self.scrollable = ScrollableFrame(self)
        # # self.scrollable.pack(fill="x", padx=10, pady=5)
        # self.scrollable = ScrollableFrame(left_frame)
        # self.scrollable.pack(fill="both", expand=True)
        # container = self.scrollable.scrollable_frame

        # --- Analysis Settings ---
        ret_frame = ttk.LabelFrame(container, text="Excerpt Retrieval Settings")
        ret_frame.pack(fill="x", padx=10, pady=5)
        ttk.Label(ret_frame, text="Score Threshold (0.0 - 1.0):").grid(row=0, column=0, padx=5, pady=5, sticky="e")
        self.score_thresh_entry = ttk.Entry(ret_frame)
        self.score_thresh_entry.insert(0, "0.3")
        self.score_thresh_entry.grid(row=0, column=1, padx=5, pady=5)
        ttk.Label(ret_frame, text="Number of Results:").grid(row=1, column=0, padx=5, pady=5, sticky="e")
        self.num_results_entry = ttk.Entry(ret_frame)
        self.num_results_entry.insert(0, "20")
        self.num_results_entry.grid(row=1, column=1, padx=5, pady=5)

        # LLM Settings
        llm_frame = ttk.LabelFrame(container, text="LLM Settings")
        llm_frame.pack(fill="x", padx=10, pady=5)
        ttk.Label(llm_frame, text="Model Selection:").grid(row=0, column=0, padx=5, pady=5, sticky="e")
        self.model_var = tk.StringVar(value="gpt4o")
        model_opts = [("gpt-4o", "gpt4o"), ("o1-mini", "o1-mini"), ("o1", "o1"), ("DeepSeek R1", "deepseek")]
        col = 1
        for lbl, val in model_opts:
            ttk.Radiobutton(llm_frame, text=lbl, variable=self.model_var, value=val).grid(row=0, column=col, padx=5, pady=5)
            col += 1
        ttk.Label(llm_frame, text="Temperature (0.0 - 1.0):").grid(row=1, column=0, padx=5, pady=5, sticky="e")
        self.temp_entry = ttk.Entry(llm_frame)
        self.temp_entry.insert(0, "0.0")
        self.temp_entry.grid(row=1, column=1, padx=5, pady=5)
        ttk.Label(llm_frame, text="Max Tokens:").grid(row=1, column=2, padx=5, pady=5, sticky="e")
        self.max_tokens_entry = ttk.Entry(llm_frame)
        self.max_tokens_entry.insert(0, "2000")
        self.max_tokens_entry.grid(row=1, column=3, padx=5, pady=5)

        # Prompt Instructions & Response Format
        prompt_frame = ttk.LabelFrame(container, text="Prompt Instructions & Response Format")
        prompt_frame.pack(fill="x", padx=10, pady=5)
        ttk.Label(prompt_frame, text="Instructions:").pack(anchor="w", padx=5, pady=2)
        self.instructions_text = scrolledtext.ScrolledText(prompt_frame, height=4)
        default_instructions = (
        "Provide a concise and specific answer in 3-5 bullet points. Each bullet point must include:\n"
        "1. A clear summary of the insight or opinion in a single sentence.\n"
        "2. A full supporting quote from the interviews, written verbatim, with no omissions or ellipses. "
        "Ensure the quote is at least 2-3 sentences long and provides enough context for the reader to "
        "understand its meaning fully.\n"
        "3. If contradictory opinions exist, provide them in a separate bullet point with a full quote.\n"
        "4. Ensure the entire output adheres to the above criteria."
    )
        self.instructions_text.insert("1.0", default_instructions)
        self.instructions_text.pack(fill="x", padx=5, pady=2)

        ttk.Label(prompt_frame, text="Response Format:").pack(anchor="w", padx=5, pady=2)
        self.response_format_text = scrolledtext.ScrolledText(prompt_frame, height=3)
        default_format = "[Summary Point]: Description of strategy or structure.\n[Quote]: From [All Metadata]"
        self.response_format_text.insert("1.0", default_format)
        self.response_format_text.pack(fill="x", padx=5, pady=2)

        # --- Grouping Options ---
        grouping_frame = ttk.LabelFrame(container, text="Data Segmentation")
        grouping_frame.pack(fill="x", padx=5, pady=5)
        ttk.Label(grouping_frame, text="Group by:").pack(side="left", padx=5, pady=5)
        self.group_by_field = tk.StringVar(value="<Select Field>")
        self.metadata_option_menu = ttk.OptionMenu(grouping_frame, self.group_by_field, "<Select Field>")
        self.metadata_option_menu.pack(side="left", padx=5, pady=5)
        ttk.Button(grouping_frame, text="Load Segment", command=self.load_divisions).pack(side="left", padx=5, pady=5)
        self.update_metadata_option_menu()

        # --- Analysis Question Section ---
        question_frame = ttk.LabelFrame(container, text="Add Question")
        question_frame.pack(fill="x", padx=5, pady=5)
        ttk.Label(question_frame, text="Question:").pack(side="left", padx=5, pady=5)
        self.new_question_entry = ttk.Entry(question_frame, width=50)
        self.new_question_entry.pack(side="left", padx=5, pady=5)
        ttk.Button(question_frame, text="Add Question Column", command=self.add_question_column).pack(side="left", padx=5, pady=5)
        generate_button = ttk.Button(question_frame, text="Fill in Grid", command=self.process_analysis)
        generate_button.pack(side="left", padx=5, pady=5)
        save_button = ttk.Button(question_frame, text="Save Grid to CSV", command=self.save_grid_to_csv)
        save_button.pack(side="left", padx=5, pady=5)

        # --- Grid view for the analysis grid ---
        grid_frame = ttk.LabelFrame(container, text="Analysis Grid")
        grid_frame.pack(fill="both", expand=True, padx=5, pady=5)
        self.tree = ttk.Treeview(grid_frame)
        self.tree.pack(side="left", fill="both", expand=True)
        scrollbar = ttk.Scrollbar(grid_frame, orient="vertical", command=self.tree.yview)
        scrollbar.pack(side="right", fill="y")
        self.tree.configure(yscrollcommand=scrollbar.set)

        self.divisions = []   # group values (rows)
        self.questions = []   # analysis question columns
        self.grid_data = {}   # {division: {question: answer}}

    def update_metadata_option_menu(self):
        """
        Refresh the metadata field list in the "Group by" dropdown.
        """
        if self.controller.all_doc_metadata:
            metadata_fields = set()
            for meta in self.controller.all_doc_metadata:
                metadata_fields.update(meta.keys())
            options = sorted(list(metadata_fields))
            if not options:
                options = ["<No metadata fields>"]
        else:
            options = ["<No metadata fields>"]
        menu = self.metadata_option_menu["menu"]
        menu.delete(0, "end")
        for opt in options:
            menu.add_command(label=opt, command=lambda value=opt: self.group_by_field.set(value))
        self.group_by_field.set(options[0])

    def load_divisions(self):
        """
        Load all distinct values (divisions) for the selected metadata field into self.divisions.
        """
        field = self.group_by_field.get()
        if field in ["<No metadata fields>", ""]:
            messagebox.showerror("Error", "No metadata field selected for grouping.")
            return

        try:
            client = self.controller.vector_store.client
        except Exception:
            try:
                from qdrant_client import QdrantClient
                client = QdrantClient(path=self.controller.qdrant_folder)
                self.controller.vector_store = QdrantVectorStore(
                    client=client,
                    collection_name="qualitative_vector_store",
                    embeddings=StanfordSecureEmbeddings()
                )
            except Exception as ex:
                messagebox.showerror("Error", f"Failed to load Qdrant: {ex}")
                return

        # Gather all points to see the unique values for that field
        all_points = []
        offset = 0
        while True:
            batch, next_page = client.scroll(collection_name="qualitative_vector_store", limit=100, offset=offset, with_payload=True, with_vectors=False)
            if not batch:
                break
            all_points.extend(batch)
            if next_page is None:
                break
            offset = next_page

        metadata_all = [p.payload.get("metadata", {}) for p in all_points]
        divisions = list({meta.get(field, "") for meta in metadata_all if meta.get(field, "")})
        self.divisions = sorted(divisions)

        # Rebuild the grid_data as blank
        self.grid_data = {}
        for div in self.divisions:
            self.grid_data[div] = {}
            # For each existing question, store blank fields
            for q in self.questions:
                self.grid_data[div][q] = {"answer": "", "excerpts": ""}

        self.refresh_grid_view()

    def refresh_grid_view(self):
        """
        Rebuild the Treeview columns and rows so that each division is a row,
        and for each question, we have two columns: 'Answer' and 'Excerpts'.
        """
        # Clear old tree
        for col in self.tree["columns"]:
            self.tree.heading(col, text="")
        self.tree.delete(*self.tree.get_children())

        # Build column list: first the grouping field, then 2 columns per question
        # e.g. if group_by_field is "Company", columns might be ["Company", "Question1 (Answer)", "Question1 (Excerpts)", ...]
        cols = [self.group_by_field.get()]
        for q in self.questions:
            cols.append(f"{q} - Answer")
            cols.append(f"{q} - Excerpts")

        self.tree["columns"] = cols[1:]  # The first column is the tree "text" column
        self.tree.heading("#0", text=cols[0])  # Use the first col as the row label

        for col in cols[1:]:
            self.tree.heading(col, text=col)
            self.tree.column(col, width=150, anchor="w")

        # Insert rows: each row is from self.divisions
        for div in self.divisions:
            # The "text" column (#0) will show the division label
            row_vals = []
            # Then we pick two values per question
            for q in self.questions:
                qa_dict = self.grid_data[div].get(q, {})
                row_vals.append(qa_dict.get("answer", ""))
                row_vals.append(qa_dict.get("excerpts", ""))
            self.tree.insert("", "end", text=div, values=row_vals)

    def add_question_column(self):
        """
        Add a new question to the grid, with blank placeholders for each division.
        """
        question = self.new_question_entry.get().strip()
        if not question:
            messagebox.showwarning("Input Error", "Please enter a question.")
            return

        # Add question to the list
        self.questions.append(question)

        # Initialize answer/excerpts for existing divisions
        for div in self.divisions:
            self.grid_data[div][question] = {"answer": "", "excerpts": ""}

        self.new_question_entry.delete(0, tk.END)
        self.refresh_grid_view()

    def process_analysis(self):
        """
        For each division and each question, retrieve relevant excerpts, pass to LLM,
        store final LLM 'answer' and 'excerpts' in self.grid_data, then refresh the grid.
        """
        try:
            retrieval_thresh = float(self.score_thresh_entry.get())
            num_results = int(self.num_results_entry.get())
            max_tokens = int(self.max_tokens_entry.get().strip())
            if max_tokens > 4000:
                messagebox.showerror("Error", "Maximum tokens cannot exceed 4000.")
                return
            temperature = float(self.temp_entry.get().strip())
            model_choice = self.model_var.get()
        except ValueError:
            messagebox.showerror("Error", "Invalid numeric entry in retrieval or token settings.")
            return

        instructions = self.instructions_text.get("1.0", tk.END).strip()
        resp_format = self.response_format_text.get("1.0", tk.END).strip()

        # Make sure we have a vector_store client
        try:
            client = self.controller.vector_store.client
        except Exception:
            client = QdrantClient(path=self.controller.qdrant_folder)
            self.controller.vector_store = QdrantVectorStore(
                client=client,
                collection_name="qualitative_vector_store",
                embeddings=StanfordSecureEmbeddings()
            )

        # Provide references to the embeddings/LLM
        embeddings = StanfordSecureEmbeddings()  # or however you define your embeddings
        if model_choice in ["gpt4o", "o1-mini", "o1"]:
            llm = StanfordSecureGPT(temperature=temperature, max_tokens=max_tokens, model_selection=model_choice)
        elif model_choice == "deepseek":
            llm = DeepSeekR1Client(temperature=temperature, max_tokens=max_tokens)
        else:
            messagebox.showerror("Error", "Unsupported model choice.")
            return

        # For progress tracking
        total = len(self.divisions) * len(self.questions)
        count = 0

        # Iterate over each division (row)
        for div in self.divisions:
            # Build a Qdrant filter for that division
            field_cond = FieldCondition(key=f"metadata.{self.group_by_field.get()}", match=MatchValue(value=div))
            q_filter = Filter(must=[field_cond])

            # For each question
            for question in self.questions:
                # 1) Embed the question and retrieve context from Qdrant
                try:
                    query_vec = embeddings.embed_query(question)
                    results = client.search(
                        collection_name="qualitative_vector_store",
                        query_vector=query_vec,
                        limit=num_results,
                        with_payload=True,
                        score_threshold=retrieval_thresh,
                        query_filter=q_filter
                    )
                except Exception as e:
                    self.grid_data[div][question] = {
                        "answer": f"Retrieval error: {e}",
                        "excerpts": ""
                    }
                    continue

                # 2) Collect the retrieved chunks into a single string
                context = ""
                for r in results:
                    meta = r.payload.get("metadata", {})
                    text = r.payload.get("page_content", "")
                    context += f"[Score: {r.score:.4f}, File: {meta.get('filename','N/A')}]\n{text}\n\n"

                # 3) Generate an LLM prompt
                final_prompt = (
                    f"Instructions:\n{instructions}\n\n"
                    f"Response Format:\n{resp_format}\n\n"
                    f"Context:\n{context}\n\n"
                    f"Question: {question}\nAnswer:"
                )

                # 4) Call the LLM
                try:
                    answer = llm.invoke(final_prompt)
                except Exception as e:
                    answer = f"LLM error: {e}"

                # 5) Store the LLM answer + the retrieved excerpts
                self.grid_data[div][question] = {
                    "answer": answer,
                    "excerpts": context
                }

                count += 1
                print(f"[AnalysisGrid] Processed {count}/{total} cells.")

        # Refresh the grid to show updated data
        self.refresh_grid_view()
        messagebox.showinfo("Analysis Complete", "Analysis grid processing complete.")

    def save_grid_to_csv(self):
        """
        Save the entire analysis grid (including 'answer' and 'excerpts' per question)
        to a CSV file picked by the user. Each row: [division, Q1-answer, Q1-excerpts, Q2-answer, Q2-excerpts, ...].
        """
        file_path = filedialog.asksaveasfilename(
            defaultextension=".csv",
            title="Save Grid as CSV",
            filetypes=[("CSV files", "*.csv")],
            initialdir=os.path.join(os.path.expanduser("~"), "Downloads"),
            initialfile="analysis_grid_results.csv"
        )
        if not file_path:
            return

        try:
            with open(file_path, "w", newline="", encoding="utf-8") as f:
                writer = csv.writer(f)
                # Build the header: first col is grouping field, then 2 columns per question
                header = [self.group_by_field.get()]
                for q in self.questions:
                    header.append(f"{q} - Answer")
                    header.append(f"{q} - Excerpts")
                writer.writerow(header)

                # Write each division as a row
                for div in self.divisions:
                    row = [div]
                    for q in self.questions:
                        qa_dict = self.grid_data[div].get(q, {})
                        ans = qa_dict.get("answer", "")
                        exc = qa_dict.get("excerpts", "")
                        row.extend([ans, exc])
                    writer.writerow(row)

            messagebox.showinfo("Success", f"Analysis grid saved to {file_path}")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save grid: {e}")

######################################################
# RUN THE APP
######################################################
if __name__ == "__main__":
    app = App()
    app.mainloop()
    app.mainloop()
